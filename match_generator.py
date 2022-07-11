# install python-docx library
# %pip install python-docx

from docx import Document
import csv

def make_sentence(name1, name2, idx, table):
  sentence = '당신은\"' + name1 + '\"입니다! \"' + name2 + '\"을(를) 찾아주세요!'
  row = table.rows[idx].cells
  row[0].text = sentence

# Load a word file
# doc = Document('drive/MyDrive/교회/file.docx')

# Make a new word file
doc= Document()
table = doc.add_table(rows = 100, cols=1)
table.style = doc.styles['Table Grid']

# open a data file(saved as csv)
f = open('drive/MyDrive/교회/data.csv', 'r', encoding='utf-8')

# read the csv file by line
rdr = csv.reader(f)

# idx stores the index of csv file , line stores the data of each line(by list type)
for idx, line in enumerate(rdr):
  # We need two sentence for each pair
  # So, the sentence index is 2*idx, 2*idx+1
  # 0,1 / 2,3 / 4,5 ...
  make_sentence(line[0], line[1], 2*idx, table)
  make_sentence(line[1], line[0], 2*idx+1, table)
f.close()

doc.save('doc_new.docx')